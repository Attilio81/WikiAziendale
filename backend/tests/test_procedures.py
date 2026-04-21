import io
import pytest

VALID_PAYLOAD = {
    "titolo": "Gestione DDT fornitore",
    "categoria": "Magazzino",
    "contenuto_md": "## Procedura\n\nRicevere il DDT e verificare la merce.",
    "autore": "Mario Rossi",
    "tags": ["magazzino", "DDT"],
}


async def test_create_procedure_returns_201(client):
    resp = await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    assert resp.status_code == 201
    data = resp.json()
    assert data["titolo"] == VALID_PAYLOAD["titolo"]
    assert data["compilation_status"] == "pending"
    assert data["version"] == 1
    assert "id" in data


async def test_create_procedure_missing_titolo_returns_422(client):
    resp = await client.post("/api/v1/procedures/", json={"contenuto_md": "x"})
    assert resp.status_code == 422


async def test_create_procedure_empty_titolo_returns_422(client):
    resp = await client.post("/api/v1/procedures/", json={"titolo": "", "contenuto_md": "x"})
    assert resp.status_code == 422


async def test_list_procedures_empty(client):
    resp = await client.get("/api/v1/procedures/")
    assert resp.status_code == 200
    data = resp.json()
    assert data["items"] == []
    assert data["total"] == 0


async def test_list_procedures_returns_created(client):
    await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    resp = await client.get("/api/v1/procedures/")
    assert resp.status_code == 200
    assert resp.json()["total"] == 1


async def test_list_procedures_filter_by_categoria(client):
    await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    await client.post("/api/v1/procedures/", json={**VALID_PAYLOAD, "titolo": "Altra", "categoria": "Qualità"})
    resp = await client.get("/api/v1/procedures/", params={"categoria": "Magazzino"})
    assert resp.json()["total"] == 1


async def test_list_procedures_search_by_title(client):
    await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    await client.post("/api/v1/procedures/", json={**VALID_PAYLOAD, "titolo": "Reso cliente"})
    resp = await client.get("/api/v1/procedures/", params={"q": "DDT"})
    data = resp.json()
    assert data["total"] == 1
    assert data["items"][0]["titolo"] == VALID_PAYLOAD["titolo"]


async def test_get_procedure_by_id(client):
    create_resp = await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    proc_id = create_resp.json()["id"]
    resp = await client.get(f"/api/v1/procedures/{proc_id}")
    assert resp.status_code == 200
    assert resp.json()["id"] == proc_id


async def test_get_procedure_not_found(client):
    resp = await client.get("/api/v1/procedures/00000000-0000-0000-0000-000000000000")
    assert resp.status_code == 404


async def test_update_procedure_increments_version(client):
    create_resp = await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    proc_id = create_resp.json()["id"]
    resp = await client.put(f"/api/v1/procedures/{proc_id}", json={"titolo": "Titolo aggiornato"})
    assert resp.status_code == 200
    data = resp.json()
    assert data["titolo"] == "Titolo aggiornato"
    assert data["version"] == 2
    assert data["compilation_status"] == "pending"


async def test_update_nonexistent_returns_404(client):
    resp = await client.put(
        "/api/v1/procedures/00000000-0000-0000-0000-000000000000",
        json={"titolo": "X"},
    )
    assert resp.status_code == 404


async def test_delete_procedure(client):
    create_resp = await client.post("/api/v1/procedures/", json=VALID_PAYLOAD)
    proc_id = create_resp.json()["id"]
    del_resp = await client.delete(f"/api/v1/procedures/{proc_id}")
    assert del_resp.status_code == 204
    get_resp = await client.get(f"/api/v1/procedures/{proc_id}")
    assert get_resp.status_code == 404


async def test_delete_nonexistent_returns_404(client):
    resp = await client.delete("/api/v1/procedures/00000000-0000-0000-0000-000000000000")
    assert resp.status_code == 404


async def test_missing_api_key_returns_error(client):
    from httpx import AsyncClient, ASGITransport
    from app.main import app as test_app
    async with AsyncClient(
        transport=ASGITransport(app=test_app),
        base_url="http://test",
    ) as ac:
        resp = await ac.get("/api/v1/procedures/")
    assert resp.status_code in (401, 422)


async def test_wrong_api_key_returns_401(client):
    from httpx import AsyncClient, ASGITransport
    from app.main import app as test_app
    async with AsyncClient(
        transport=ASGITransport(app=test_app),
        base_url="http://test",
        headers={"X-API-Key": "wrong-key"},
    ) as ac:
        resp = await ac.get("/api/v1/procedures/")
    assert resp.status_code == 401


async def test_create_procedure_triggers_compilation(client):
    """POST /procedures returns 201 — background task is registered but not awaited in test."""
    payload = {
        "titolo": "Nuova Procedura",
        "categoria": "Test",
        "contenuto_md": "## Test\n\nContenuto.",
        "autore": "Test",
        "tags": [],
    }
    r = await client.post("/api/v1/procedures/", json=payload)
    assert r.status_code == 201
    # compilation_status starts as "pending"; background task runs after response
    assert r.json()["compilation_status"] == "pending"


async def test_update_procedure_resets_compilation_status(client):
    """PUT /procedures/{id} resets compilation_status to pending."""
    # Create
    payload = {
        "titolo": "Procedura Da Aggiornare",
        "categoria": "Test",
        "contenuto_md": "## Originale",
        "autore": None,
        "tags": [],
    }
    r = await client.post("/api/v1/procedures/", json=payload)
    proc_id = r.json()["id"]

    # Update
    r = await client.put(f"/api/v1/procedures/{proc_id}", json={"titolo": "Aggiornata"})
    assert r.status_code == 200
    assert r.json()["compilation_status"] == "pending"


# --- Upload helpers ---

def _make_docx(text: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_txt():
    from app.api.procedures import _extract_txt
    result = _extract_txt(b"Procedura aziendale\nPasso 1: fare qualcosa")
    assert "Procedura aziendale" in result
    assert "Passo 1" in result


def test_extract_docx():
    from app.api.procedures import _extract_docx
    result = _extract_docx(_make_docx("Contenuto procedura di test"))
    assert "Contenuto procedura di test" in result


def test_extract_docx_multiple_paragraphs():
    from app.api.procedures import _extract_docx
    from docx import Document
    doc = Document()
    doc.add_paragraph("Primo paragrafo")
    doc.add_paragraph("Secondo paragrafo")
    buf = io.BytesIO()
    doc.save(buf)
    result = _extract_docx(buf.getvalue())
    assert "Primo paragrafo" in result
    assert "Secondo paragrafo" in result


def test_derive_title_from_filename():
    from app.api.procedures import _derive_title
    assert _derive_title("ricezione-merci.pdf") == "Ricezione Merci"
    assert _derive_title("gestione_non_conformita.docx") == "Gestione Non Conformita"
    assert _derive_title("Procedura Acquisti.txt") == "Procedura Acquisti"


# --- Upload endpoint ---

async def test_upload_txt_creates_procedure(client):
    content = b"Procedura di magazzino\n\nPasso 1: controllare DDT\nPasso 2: firmare registro"
    resp = await client.post(
        "/api/v1/procedures/upload",
        files={"file": ("ricezione-merci.txt", content, "text/plain")},
    )
    assert resp.status_code == 201
    data = resp.json()
    assert data["titolo"] == "Ricezione Merci"
    assert data["compilation_status"] == "pending"
    assert "DDT" in data["contenuto_md"]


async def test_upload_docx_creates_procedure(client):
    docx_bytes = _make_docx("Procedura di qualità per la gestione delle non conformità")
    resp = await client.post(
        "/api/v1/procedures/upload",
        files={"file": ("gestione-nc.docx", docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 201
    data = resp.json()
    assert data["titolo"] == "Gestione Nc"
    assert data["compilation_status"] == "pending"


async def test_upload_with_categoria_and_autore(client):
    resp = await client.post(
        "/api/v1/procedures/upload",
        files={"file": ("test.txt", b"Procedura di test", "text/plain")},
        data={"categoria": "Qualità", "autore": "Mario Rossi"},
    )
    assert resp.status_code == 201
    data = resp.json()
    assert data["categoria"] == "Qualità"
    assert data["autore"] == "Mario Rossi"


async def test_upload_unsupported_format_returns_400(client):
    resp = await client.post(
        "/api/v1/procedures/upload",
        files={"file": ("document.xlsx", b"fake", "application/octet-stream")},
    )
    assert resp.status_code == 400
    assert "non supportato" in resp.json()["detail"]


async def test_upload_empty_content_returns_400(client):
    resp = await client.post(
        "/api/v1/procedures/upload",
        files={"file": ("vuoto.txt", b"   \n\n   ", "text/plain")},
    )
    assert resp.status_code == 400
    assert "estrarre testo" in resp.json()["detail"]


async def test_upload_no_api_key_returns_401(client):
    from httpx import AsyncClient, ASGITransport
    from app.main import app as test_app
    async with AsyncClient(transport=ASGITransport(app=test_app), base_url="http://test") as ac:
        resp = await ac.post(
            "/api/v1/procedures/upload",
            files={"file": ("test.txt", b"contenuto", "text/plain")},
        )
    assert resp.status_code == 401
